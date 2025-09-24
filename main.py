import os
from openai import OpenAI
from dotenv import load_dotenv
from moviepy.editor import VideoFileClip
from pydantic import BaseModel, Field
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import PydanticOutputParser
from langchain.schema.messages import HumanMessage
from docx import Document
from typing import Optional, List
import tempfile
from moviepy.editor import AudioFileClip

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class HealthHistory(BaseModel):
    age: Optional[int] = Field(None, description="Age of the client")
    gender: Optional[str] = Field(None, description="Gender of the client")
    marital_status: Optional[str] = Field(None, description="Marital Status of the client")
    city: Optional[str] = Field(None, description="City of the client")
    food_preference: Optional[str] = Field(None, description="Food Preference of the client")
    height_feet: Optional[str] = Field(None, description="Height in Feet")
    weight_kg: Optional[str] = Field(None, description="Weight in KGs")
    profession: Optional[str] = Field(None, description="Profession of the client")

    health_history: Optional[List[str]] = Field(description="Health history of the client")
    discussion_recommendations: Optional[List[str]] = Field(description="Recommendations and discussion")
    warm_message: Optional[str] = Field(description="Warm message from the consultant")

parser = PydanticOutputParser(pydantic_object=HealthHistory)
format_instructions = parser.get_format_instructions()

def extract_audio(file_path: str) -> str:
    """
    Extract audio from video file. Handles both video + audio MP4 and audio-only MP4.
    Returns path to temp .mp3 file.
    """
    temp_audio_path = tempfile.mktemp(suffix=".mp3")

    try:
        # Try as video first
        with VideoFileClip(file_path) as video:
            if video.audio is None:
                raise ValueError("No audio stream found in video file")
            video.audio.write_audiofile(temp_audio_path)
    except Exception:
        # If failed (likely audio-only MP4), fallback to AudioFileClip
        with AudioFileClip(file_path) as audio:
            audio.write_audiofile(temp_audio_path)

    return temp_audio_path

def split_audio(audio_path: str, chunk_length_sec: int = 540) -> List[str]:
    audio = AudioFileClip(audio_path)
    duration = int(audio.duration)
    chunk_paths = []

    for i in range(0, duration, chunk_length_sec):
        start = i
        end = min(i + chunk_length_sec, duration)
        chunk = audio.subclip(start, end)

        chunk_path = tempfile.mktemp(suffix=f"_chunk_{i//chunk_length_sec}.mp3")
        chunk.write_audiofile(chunk_path)
        chunk_paths.append(chunk_path)

    audio.close()
    return chunk_paths

def transcribe_audio(audio_path: str) -> str:
    chunks = split_audio(audio_path)
    full_transcript = ""
    for i, chunk_path in enumerate(chunks):
        with open(chunk_path, "rb") as audio_file:
            print(f"Transcribing chunk {i+1}/{len(chunks)}...")
            response = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="text"
            )
            full_transcript += response + "\n"
        os.remove(chunk_path)  # Clean up each chunk
    return full_transcript.strip()

def generate_summary(transcript: str, preconsult: str) -> HealthHistory:
    PROMPT = f"""
    The following text is a transcript of a medical consultation that was originally recorded as an audio file. This transcript includes the client’s medical background and a consultation discussion with Luke Coutinho.

    Extract information for these categories:
    - Age, Gender, Marital Status, City, Food Preference, Height in Feet, Weight in KGs, Profession

    Your task is to extract two separate sections from this audio text:

    1. **Health History**:
    - Extract all relevant details about the client’s existing or past health conditions (e.g., diabetes, hypertension, digestive issues, etc.)
    - Include the duration if mentioned (e.g., "for 10 years")
    - Also include any family history, lifestyle habits, medications, or supplements the client is currently taking
    - Present this as clear points

    2. **Discussion & Recommendations by Luke**:
    - Summarize what Luke discussed with the client regarding their condition
    - Mention any advice given, including lifestyle or dietary suggestions
    - Include any recommended supplements, tests, or follow-up actions
    - Present this section in point format as well

    If any section is missing or not clearly mentioned, just write “Not mentioned.”

    3. **Warm Message**:
    The tone should reflect Luke Coutinho style — warm, practical, honest, and motivating. Speak to the heart. Do NOT include any headings or labels beyond what is requested.

    Return the response in this exact JSON format:
    {format_instructions}
    """

    llm = ChatOpenAI(model="gpt-4.1-mini", temperature=0.2)
    response = llm.invoke([HumanMessage(content=PROMPT + "\n\n" + preconsult + "\n\n" + transcript)])
    return parser.parse(response.content)

def extract_text_from_word_filelike(file) -> str:
    doc = Document(file)
    texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                texts.append(" | ".join(row_data))
    return "\n".join(texts)

def summarize_consultation(file_path: str, pre_extracted: str) -> dict:
    audio_path = None
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        video_extensions = {'.mp4', '.mov', '.avi', '.mkv'}
        audio_extensions = {'.mp3', '.wav', '.m4a', '.aac', '.ogg','.mp4'}
        
        if file_ext in video_extensions:
            audio_path = extract_audio(file_path)
            transcript = transcribe_audio(audio_path)
        elif file_ext in audio_extensions:
            transcript = transcribe_audio(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        summary = generate_summary(transcript, pre_extracted)
        print(summary.dict())  # Debugging line
        return summary.dict()

    finally:
        if audio_path and os.path.exists(audio_path):
            os.remove(audio_path)

