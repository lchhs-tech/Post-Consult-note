from docx import Document
import tempfile

def write_summary_to_docx(summaries: list) -> str:
    doc = Document()
    doc.add_heading("Post Consultation Summary", level=1)

    if summaries:

        for idx, item in enumerate(summaries, 1):

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(
                f"{item.get('age', 'N/A')}/ "
                f"{item.get('gender', 'N/A')}/ "
                f"{item.get('marital_status', 'N/A')}/ "
                f"{item.get('city', 'N/A')}/ "
                f"{item.get('food_preference', 'N/A')}/ "
                f"{item.get('height_feet', 'N/A')}/ "
                f"{item.get('weight_kg', 'N/A')}/ "
                f"{item.get('profession', 'N/A')}"
            )
            run.bold = True

        # Health History Section
        doc.add_heading("Health History", level=3)
        if item.get("health_history"):
            for line in item.get("health_history", []):
                doc.add_paragraph(line, style="List Bullet")

        # Discussion & Recommendations Section
        doc.add_heading("Discussion & Recommendations", level=3)
        if item.get("discussion_recommendations"):
            for point in item.get("discussion_recommendations", []):
                doc.add_paragraph(point, style="List Bullet")

        
        if item.get("warm_message"):
            doc.add_paragraph(item.get("warm_message", "Not mentioned"))    
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
    return tmp.name